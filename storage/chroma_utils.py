import os
import tempfile
from dotenv import load_dotenv
import fitz
import docx
from langchain_openai import AzureOpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

load_dotenv()

embedding_endpoint = os.getenv("EMBEDDING_ENDPOINT")
embedding_api = os.getenv("EMBEDDING_API_KEY")
embeddings = AzureOpenAIEmbeddings(
    azure_deployment="text-embedding-3-large",
    openai_api_version="2023-05-15",
    azure_endpoint=embedding_endpoint,
    api_key=embedding_api,
)

CHROMA_DB_PATH = "./chroma_db"


def extract_text_from_pdf(uploaded_file):
    try:
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        all_text = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            all_text.append(f"--- Page {page_num + 1} ---\n{text}")
        return "\n".join(all_text)
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return ""


def extract_text_from_docx(uploaded_file):
    try:
        temp = tempfile.NamedTemporaryFile(delete=False)
        temp.write(uploaded_file.read())
        temp.close()
        doc = docx.Document(temp.name)
        full_text = [para.text for para in doc.paragraphs]
        os.unlink(temp.name)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return ""


def extract_text_from_txt(uploaded_file):
    try:
        return uploaded_file.read().decode("utf-8")
    except Exception as e:
        print(f"Error extracting text from TXT: {str(e)}")
        return ""


def is_youtube_context_file(text, filename):
    """
    Detect if this is a YouTube context file based on content patterns.
    """
    indicators = [
        "VIDEO_" in text,
        "TITLE:" in text,
        "TOPICS:" in text,
        "CONTENT:" in text,
        "youtube.com/watch" in text.lower(),
        "youtube" in filename.lower(),
    ]
    return sum(indicators) >= 3


def smart_chunk_youtube_content(text):
    """
    Intelligently chunk YouTube context files by video entries.
    Each chunk = 1 complete video with all its metadata.
    """
    chunks = []

    # Split by the video separator pattern
    if "VIDEO_" in text:
        # Split by VIDEO_N markers
        video_entries = text.split("VIDEO_")

        for entry in video_entries:
            entry = entry.strip()
            if not entry:
                continue

            # Reconstruct with VIDEO_ prefix
            full_entry = f"VIDEO_{entry}"

            # Each video is a complete chunk
            if len(full_entry) > 50:  # Skip very short entries
                chunks.append(full_entry)

    # Fallback: if no VIDEO_ markers found, try other separators
    elif "Title:" in text and "URL:" in text:
        # Old format with Title/URL/Description
        entries = text.split(
            "--------------------------------------------------------------------------------"
        )

        for entry in entries:
            entry = entry.strip()
            if entry and len(entry) > 50 and ("Title:" in entry or "URL:" in entry):
                chunks.append(entry)

    else:
        # Not a YouTube file, return None to use default chunking
        return None

    print(f"Smart chunking created {len(chunks)} video-based chunks")
    return chunks


def get_text_splitter(text, filename):
    """
    Return appropriate text splitter based on content type.
    """
    # Check if it's a YouTube context file
    if is_youtube_context_file(text, filename):
        print(f"Detected YouTube context file: {filename}")
        # Return None to signal we'll use smart chunking instead
        return None

    # Default splitter for regular documents
    return RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=200, separators=["\n\n", "\n", " ", ""]
    )


def delete_file_from_chroma(chatbot_id, filename):
    """
    Delete all chunks belonging to a specific filename from ChromaDB.
    Returns number of chunks deleted.
    """
    try:
        collection_name = f"chatbot_{chatbot_id}"
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=embeddings,
            persist_directory=CHROMA_DB_PATH,
        )

        # Get all documents with their metadata
        collection = vector_store._collection
        existing_docs = collection.get(include=["metadatas"])

        if not existing_docs["metadatas"]:
            print(f"No documents found in ChromaDB for chatbot {chatbot_id}")
            return 0

        # Find chunks matching the filename (without extension)
        filename_base = os.path.splitext(filename)[0]
        doc_ids_to_delete = []

        for idx, metadata in enumerate(existing_docs["metadatas"]):
            if metadata.get("filename") == filename_base:
                doc_ids_to_delete.append(existing_docs["ids"][idx])

        if doc_ids_to_delete:
            collection.delete(ids=doc_ids_to_delete)
            print(
                f"✅ Deleted {len(doc_ids_to_delete)} chunks for file '{filename_base}'"
            )
            return len(doc_ids_to_delete)
        else:
            print(f"⚠️ No chunks found for file '{filename_base}' in ChromaDB")
            return 0

    except Exception as e:
        print(f"❌ Error deleting file from ChromaDB: {e}")
        return 0


def update_consolidated_content(files_data, chatbot_id, chatbot_name):
    """
    Updates the chatbot's documents in ChromaDB.
    IMPORTANT: Only replaces files that are being re-uploaded, keeps other files intact.
    """
    try:
        collection_name = f"chatbot_{chatbot_id}"
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=embeddings,
            persist_directory=CHROMA_DB_PATH,
        )

        consolidated_file = f"chatbot_{chatbot_id}_consolidated.txt"
        all_text = ""
        documents = []

        # Track existing count before changes
        existing_count = 0
        try:
            existing_count = vector_store._collection.count()
        except Exception:
            existing_count = 0

        # STEP 1: Only delete files that are being RE-UPLOADED (replace existing versions)
        # Check which files already exist in ChromaDB
        collection = vector_store._collection
        existing_docs = collection.get(include=["metadatas"])
        existing_filenames = set()
        if existing_docs["metadatas"]:
            existing_filenames = {
                doc.get("filename") for doc in existing_docs["metadatas"]
            }

        deleted_count = 0
        for file_data in files_data:
            filename_base = os.path.splitext(file_data["filename"])[0]
            # Only delete if this file already exists (we're replacing it)
            if filename_base in existing_filenames:
                deleted = delete_file_from_chroma(chatbot_id, file_data["filename"])
                deleted_count += deleted
                print(f"🔄 Replacing existing file: {filename_base}")
            else:
                print(f"➕ Adding new file: {filename_base}")

        if deleted_count > 0:
            print(f"📝 Deleted {deleted_count} old chunks from replaced files")

        # STEP 2: Process and add new file versions
        for doc_id, file_data in enumerate(files_data, start=1):
            uploaded_file = file_data["file"]
            filename = os.path.splitext(file_data["filename"])[0]
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()

            # Extract text safely
            if file_extension == ".pdf":
                text = extract_text_from_pdf(uploaded_file) or ""
            elif file_extension == ".docx":
                text = extract_text_from_docx(uploaded_file) or ""
            elif file_extension == ".txt":
                text = extract_text_from_txt(uploaded_file) or ""
            else:
                print(f"Unsupported file type: {uploaded_file.name}")
                continue

            if not text.strip():
                text = f"Error extracting content from {filename}"

            all_text += f"\n\n--- File: {filename} ---\n\n{text}"

            # Smart chunking logic
            text_splitter = get_text_splitter(text, filename)
            if text_splitter is None:
                chunks = smart_chunk_youtube_content(text)
                if chunks is None:
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=1000, chunk_overlap=200
                    )
                    chunks = text_splitter.split_text(text)
            else:
                chunks = text_splitter.split_text(text)

            for chunk_id, chunk in enumerate(chunks):
                if not chunk or not isinstance(chunk, str):
                    continue

                metadata = {
                    "id": f"{chatbot_id}-{doc_id}-{chunk_id}",
                    "filename": filename,
                    "page_number": doc_id,
                    "chunk_id": chunk_id,
                }

                # Add YouTube video-related metadata if available
                if "Title:" in chunk or "TITLE:" in chunk:
                    for line in chunk.split("\n"):
                        if "TITLE:" in line.upper():
                            metadata["video_title"] = line.split(":", 1)[-1].strip()
                            break

                document = Document(page_content=chunk.strip(), metadata=metadata)
                documents.append(document)

            uploaded_file.seek(0)

        # STEP 3: Add new documents
        if documents:
            print(f"➕ Adding {len(documents)} new chunks to ChromaDB...")
            vector_store.add_documents(documents)
        else:
            print("⚠️ No valid documents found to add.")

        # STEP 4: Update consolidated file (overwrite with current state)
        # Read existing content first
        existing_content = ""
        if os.path.exists(consolidated_file):
            with open(consolidated_file, "r", encoding="utf-8") as f:
                existing_content = f.read()

        # Remove old file sections and add new ones
        for file_data in files_data:
            filename_base = os.path.splitext(file_data["filename"])[0]
            existing_content = remove_document_content(existing_content, filename_base)

        # Write updated content
        with open(consolidated_file, "w", encoding="utf-8") as f:
            f.write(existing_content)
            f.write(all_text)

        new_count = existing_count - deleted_count + len(documents)
        print(f"✅ Updated ChromaDB: {existing_count} → {new_count} chunks")
        print(f"   (Deleted: {deleted_count}, Added: {len(documents)})")

        return f"Uploaded {len(documents)} documents successfully"

    except Exception as e:
        print(f"❌ Error updating content: {e}")
        import traceback

        traceback.print_exc()
        return f"failed: {str(e)}"


def process_and_store_files(files_data, chatbot_id):
    """
    Creates a new Chroma collection if it doesn't exist, or adds to it incrementally.
    """
    try:
        collection_name = f"chatbot_{chatbot_id}"
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=embeddings,
            persist_directory=CHROMA_DB_PATH,
        )

        all_text = ""
        documents = []

        for doc_id, file_data in enumerate(files_data, start=1):
            uploaded_file = file_data["file"]
            filename = os.path.splitext(file_data["filename"])[0]
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()

            if file_extension == ".pdf":
                print(f"Processing PDF: {filename}")
                text = extract_text_from_pdf(uploaded_file)
            elif file_extension == ".docx":
                print(f"Processing DOCX: {filename}")
                text = extract_text_from_docx(uploaded_file)
            elif file_extension == ".txt":
                print(f"Processing TXT: {filename}")
                text = extract_text_from_txt(uploaded_file)
            else:
                print(f"Unsupported file type: {filename}{file_extension}")
                continue

            if not text:
                text = f"Error extracting content from {filename}"

            all_text += f"\n\n--- File: {filename} ---\n\n{text}"

            # Smart chunking
            text_splitter = get_text_splitter(text, filename)
            if text_splitter is None:
                chunks = smart_chunk_youtube_content(text)
                if chunks is None:
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=1000, chunk_overlap=200
                    )
                    chunks = text_splitter.split_text(text)
            else:
                chunks = text_splitter.split_text(text)

            for chunk_id, chunk in enumerate(chunks):
                if not chunk or not isinstance(chunk, str):
                    continue

                metadata = {
                    "id": f"{chatbot_id}-{doc_id}-{chunk_id}",
                    "filename": filename,
                    "page_number": doc_id,
                    "chunk_id": chunk_id,
                }

                if "Title:" in chunk or "TITLE:" in chunk:
                    for line in chunk.split("\n"):
                        if "TITLE:" in line.upper():
                            metadata["video_title"] = line.split(":", 1)[-1].strip()
                            break

                document = Document(page_content=chunk.strip(), metadata=metadata)
                documents.append(document)

            uploaded_file.seek(0)

        if documents:
            vector_store.add_documents(documents)
        else:
            print("⚠️ No valid documents extracted.")

        consolidated_file = f"chatbot_{chatbot_id}_consolidated.txt"
        with open(consolidated_file, "w", encoding="utf-8") as f:
            f.write(all_text)

        print(f"✅ Added {len(documents)} chunks to '{collection_name}'")
        return f"Uploaded {len(documents)} documents successfully"

    except Exception as e:
        print(f"❌ Error processing files: {str(e)}")
        return f"failed: {str(e)}"


def delete_document_from_chroma(chatbot_id, filename=None):
    try:
        collection_name = f"chatbot_{chatbot_id}"
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=embeddings,
            persist_directory=CHROMA_DB_PATH,
        )
        collection = vector_store._collection
        existing_docs = collection.get(include=["metadatas"])

        if not existing_docs["metadatas"]:
            return False, "No documents found in ChromaDB for this chatbot"

        if filename:
            filename = os.path.splitext(filename)[0]
            doc_ids_to_delete = [
                existing_docs["ids"][i]
                for i, doc in enumerate(existing_docs["metadatas"])
                if doc.get("filename") == filename
            ]
            if not doc_ids_to_delete:
                return False, f"No documents found with filename: {filename}"
            collection.delete(ids=doc_ids_to_delete)

            # Also remove from consolidated file
            consolidated_file = f"chatbot_{chatbot_id}_consolidated.txt"
            if os.path.exists(consolidated_file):
                with open(consolidated_file, "r", encoding="utf-8") as f:
                    content = f.read()
                updated_content = remove_document_content(content, filename)
                with open(consolidated_file, "w", encoding="utf-8") as f:
                    f.write(updated_content)

            return (
                True,
                f"Successfully deleted {len(doc_ids_to_delete)} document(s) with filename: {filename}",
            )
        else:
            vector_store.delete_collection()
            vector_store = Chroma(
                collection_name=collection_name,
                embedding_function=embeddings,
                persist_directory=CHROMA_DB_PATH,
            )
            # Also delete consolidated file
            consolidated_file = f"chatbot_{chatbot_id}_consolidated.txt"
            if os.path.exists(consolidated_file):
                os.remove(consolidated_file)
            return True, "All documents deleted successfully"
    except Exception as e:
        return False, f"Error deleting from ChromaDB: {str(e)}"


def get_document_from_chroma(chatbot_id, filename):
    try:
        collection_name = f"chatbot_{chatbot_id}"
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=embeddings,
            persist_directory=CHROMA_DB_PATH,
        )
        results = vector_store._collection.get(include=["documents", "metadatas"])
        filename = os.path.splitext(filename)[0]

        # Collect all chunks for the filename
        chunks = []
        for doc_content, metadata in zip(results["documents"], results["metadatas"]):
            if metadata.get("filename") == filename:
                chunks.append((metadata["chunk_id"], doc_content))

        if not chunks:
            return ""

        # Sort by chunk_id and concatenate
        chunks.sort(key=lambda x: x[0])  # Sort by chunk_id
        full_content = "\n".join(chunk[1] for chunk in chunks)
        return full_content
    except Exception as e:
        print(f"Error retrieving from ChromaDB: {str(e)}")
        return ""


def remove_document_content(content, filename):
    """
    Remove a specific file's content from the consolidated text.
    """
    if not content:
        return ""

    filename = os.path.splitext(filename)[0]
    sections = content.split("\n\n")

    # Filter out sections that belong to this file
    filtered_sections = []
    skip_until_next_file = False

    for section in sections:
        section_stripped = section.strip()

        # Check if this is a file header
        if section_stripped.startswith("--- File:"):
            # Check if it's the file we want to remove
            if f"--- File: {filename} ---" in section_stripped:
                skip_until_next_file = True
                continue
            else:
                skip_until_next_file = False

        # Add section if we're not skipping
        if not skip_until_next_file and section_stripped:
            filtered_sections.append(section)

    return "\n\n".join(filtered_sections).strip()
